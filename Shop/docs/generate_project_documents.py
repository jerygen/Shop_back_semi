from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].startswith("|"):
        line = lines[i].strip()
        if set(line.replace("|", "").strip()) <= {"-", ":"}:
            i += 1
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    for line in lines:
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)


def markdown_to_docx(md_path, docx_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, "000000"),
        ("Heading 2", 16, "000000"),
        ("Heading 3", 14, "434343"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    first_heading = True

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.space_after = Pt(3)
            if not first_heading:
                run.add_break(WD_BREAK.PAGE)
            first_heading = False
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
        elif line[0:3].replace(".", "").isdigit() and ". " in line[:5]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(line.split(". ", 1)[1].strip())
        else:
            doc.add_paragraph(line)
        i += 1

    doc.save(docx_path)


def main():
    for filename in ["쇼핑몰_기획_문서.md", "쇼핑몰_구현_계획서.md"]:
        markdown_to_docx(BASE_DIR / filename, BASE_DIR / filename.replace(".md", ".docx"))


if __name__ == "__main__":
    main()
